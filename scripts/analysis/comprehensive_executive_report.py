#!/usr/bin/env python3
"""
Comprehensive Executive Report Generator
Creates a professional Word document consolidating all analysis insights
"""

import json
import os
import re
from datetime import datetime
from collections import defaultdict
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np
import pandas as pd

class ComprehensiveExecutiveReport:
    def __init__(self):
        self.output_file = "reports/executive/comprehensive_executive_report_20250725.docx"
        self.data_file = "reports/executive/cleaned_licensing_data_20250725.json"
        self.industry_analysis_file = "reports/executive/enhanced_industry_analysis_20250725.json"
        self.msp_analysis_file = "reports/executive/granular_msp_analysis_20250725.json"
        self.data_quality_file = "reports/executive/data_quality_analysis_20250725.json"
        
    def load_data(self):
        """Load all analysis data."""
        data = {}
        
        # Load cleaned data
        if os.path.exists(self.data_file):
            with open(self.data_file, 'r') as f:
                data['cleaned_data'] = json.load(f)
        
        # Load industry analysis
        if os.path.exists(self.industry_analysis_file):
            with open(self.industry_analysis_file, 'r') as f:
                data['industry_analysis'] = json.load(f)
        
        # Load MSP analysis
        if os.path.exists(self.msp_analysis_file):
            with open(self.msp_analysis_file, 'r') as f:
                data['msp_analysis'] = json.load(f)
        
        # Load data quality analysis
        if os.path.exists(self.data_quality_file):
            with open(self.data_quality_file, 'r') as f:
                data['data_quality'] = json.load(f)
        
        return data
    
    def create_charts(self, data):
        """Create professional charts for the report."""
        charts_dir = "reports/charts"
        os.makedirs(charts_dir, exist_ok=True)
        
        charts = {}
        
        if 'cleaned_data' in data:
            # Spending by Category Chart
            category_spend = defaultdict(float)
            vendor_spend = defaultdict(float)
            company_spend = defaultdict(float)
            
            for item in data['cleaned_data']:
                vendor = item.get('vendor', 'Unknown')
                company = item.get('company', 'Unknown Company')
                amount = item.get('total_amount', 0)
                
                # Categorize vendor
                category = self.categorize_vendor(vendor)
                category_spend[category] += amount
                vendor_spend[vendor] += amount
                company_spend[company] += amount
            
            # Category spending pie chart
            if category_spend:
                fig, ax = plt.subplots(figsize=(10, 8))
                categories = list(category_spend.keys())
                amounts = list(category_spend.values())
                colors = ['#FF6B6B', '#4ECDC4', '#45B7D1', '#96CEB4', '#FFEAA7', '#DDA0DD']
                
                wedges, texts, autotexts = ax.pie(amounts, labels=categories, autopct='%1.1f%%', 
                                                 colors=colors[:len(categories)], startangle=90)
                ax.set_title('Spending Distribution by Category', fontsize=16, fontweight='bold')
                plt.tight_layout()
                chart_path = f"{charts_dir}/category_spending_pie.png"
                plt.savefig(chart_path, dpi=300, bbox_inches='tight')
                plt.close()
                charts['category_pie'] = chart_path
            
            # Top vendors bar chart
            if vendor_spend:
                top_vendors = sorted(vendor_spend.items(), key=lambda x: x[1], reverse=True)[:8]
                vendors, amounts = zip(*top_vendors)
                
                fig, ax = plt.subplots(figsize=(12, 8))
                bars = ax.bar(vendors, amounts, color='#45B7D1')
                ax.set_title('Top Vendors by Spend (Consolidated)', fontsize=16, fontweight='bold')
                ax.set_xlabel('Vendors', fontsize=12)
                ax.set_ylabel('Total Spend ($)', fontsize=12)
                plt.xticks(rotation=45, ha='right')
                
                # Add value labels on bars
                for bar, amount in zip(bars, amounts):
                    height = bar.get_height()
                    ax.text(bar.get_x() + bar.get_width()/2., height + height*0.01,
                           f'${amount:,.0f}', ha='center', va='bottom', fontweight='bold')
                
                plt.tight_layout()
                chart_path = f"{charts_dir}/top_vendors_bar.png"
                plt.savefig(chart_path, dpi=300, bbox_inches='tight')
                plt.close()
                charts['vendors_bar'] = chart_path
            
            # Company spending bar chart
            if company_spend:
                top_companies = sorted(company_spend.items(), key=lambda x: x[1], reverse=True)[:5]
                companies, amounts = zip(*top_companies)
                
                fig, ax = plt.subplots(figsize=(10, 6))
                bars = ax.bar(companies, amounts, color='#FF6B6B')
                ax.set_title('Spending by Company', fontsize=16, fontweight='bold')
                ax.set_xlabel('Companies', fontsize=12)
                ax.set_ylabel('Total Spend ($)', fontsize=12)
                plt.xticks(rotation=45, ha='right')
                
                # Add value labels on bars
                for bar, amount in zip(bars, amounts):
                    height = bar.get_height()
                    ax.text(bar.get_x() + bar.get_width()/2., height + height*0.01,
                           f'${amount:,.0f}', ha='center', va='bottom', fontweight='bold')
                
                plt.tight_layout()
                chart_path = f"{charts_dir}/company_spending_bar.png"
                plt.savefig(chart_path, dpi=300, bbox_inches='tight')
                plt.close()
                charts['companies_bar'] = chart_path
        
        return charts
    
    def categorize_vendor(self, vendor_name):
        """Categorize vendor based on name."""
        vendor_lower = vendor_name.lower()
        
        vendor_categories = {
            "synoptek": "it_services",
            "atlassian": "development_tools",
            "microsoft": "enterprise_software",
            "oracle": "enterprise_software",
            "salesforce": "enterprise_software",
            "aws": "cloud_services",
            "amazon": "cloud_services",
            "azure": "cloud_services",
            "google": "cloud_services",
            "gcp": "cloud_services",
            "github": "development_tools",
            "gitlab": "development_tools",
            "crowdstrike": "security_software",
            "sentinelone": "security_software",
            "palo alto": "security_software",
            "proofpoint": "security_software",
            "harman": "it_services",
            "markov": "it_services"
        }
        
        for vendor_key, category in vendor_categories.items():
            if vendor_key in vendor_lower:
                return category
        
        return "it_services"
    
    def create_executive_summary(self, data):
        """Create the executive summary section."""
        if 'cleaned_data' not in data:
            return "No data available for analysis."
        
        cleaned_data = data['cleaned_data']
        total_spend = sum(item.get('total_amount', 0) for item in cleaned_data)
        total_invoices = len(cleaned_data)
        
        # Calculate vendor and company counts
        vendors = set(item.get('vendor', 'Unknown') for item in cleaned_data)
        companies = set(item.get('company', 'Unknown Company') for item in cleaned_data)
        
        # Calculate category breakdown
        category_spend = defaultdict(float)
        for item in cleaned_data:
            vendor = item.get('vendor', 'Unknown')
            category = self.categorize_vendor(vendor)
            category_spend[category] += item.get('total_amount', 0)
        
        # Find top vendor and company
        vendor_spend = defaultdict(float)
        company_spend = defaultdict(float)
        for item in cleaned_data:
            vendor = item.get('vendor', 'Unknown')
            company = item.get('company', 'Unknown Company')
            amount = item.get('total_amount', 0)
            vendor_spend[vendor] += amount
            company_spend[company] += amount
        
        top_vendor = max(vendor_spend.items(), key=lambda x: x[1]) if vendor_spend else ("None", 0)
        top_company = max(company_spend.items(), key=lambda x: x[1]) if company_spend else ("None", 0)
        
        summary = f"""
EXECUTIVE SUMMARY

Total Annual Spend: ${total_spend:,.2f}
Total Invoices: {total_invoices:,}
Unique Vendors: {len(vendors)}
Unique Companies: {len(companies)}

Top Spending Vendor: {top_vendor[0]} (${top_vendor[1]:,.2f})
Top Spending Company: {top_company[0]} (${top_company[1]:,.2f})

Category Breakdown:
"""
        
        for category, spend in sorted(category_spend.items(), key=lambda x: x[1], reverse=True):
            percentage = (spend / total_spend) * 100
            summary += f"• {category.replace('_', ' ').title()}: ${spend:,.2f} ({percentage:.1f}%)\n"
        
        return summary
    
    def create_industry_benchmark_analysis(self, data):
        """Create industry benchmark analysis section."""
        if 'industry_analysis' not in data:
            return "Industry benchmark analysis not available."
        
        analysis = data['industry_analysis']
        summary = analysis.get('summary', {})
        benchmarks = analysis.get('benchmarks', {})
        
        section = f"""
INDUSTRY BENCHMARK ANALYSIS

Total Spend Analyzed: ${summary.get('total_spend', 0):,.2f}
Years Analyzed: {', '.join(map(str, summary.get('years_analyzed', [])))}
Vendors Consolidated: {summary.get('vendor_count', 0)}
Companies Identified: {summary.get('company_count', 0)}

Benchmark Comparison:
"""
        
        for category, benchmark_data in benchmarks.items():
            actual = benchmark_data.get('actual_spend', 0)
            variance = benchmark_data.get('variance_percentage', 0)
            status = benchmark_data.get('status', 'Unknown')
            
            section += f"• {category.replace('_', ' ').title()}: ${actual:,.2f} ({variance:+.1f}% vs benchmark) - {status}\n"
        
        # Add recommendations
        recommendations = analysis.get('recommendations', [])
        if recommendations:
            section += "\nKey Recommendations:\n"
            for i, rec in enumerate(recommendations[:5], 1):  # Top 5 recommendations
                section += f"{i}. {rec.get('message', 'No message')}\n"
        
        return section
    
    def create_msp_analysis(self, data):
        """Create MSP analysis section."""
        if 'msp_analysis' not in data:
            return "MSP analysis not available."
        
        analysis = data['msp_analysis']
        summary = analysis.get('summary', {})
        msp_services = analysis.get('msp_services', {})
        
        section = f"""
MSP SERVICE ANALYSIS

Total MSP Spend: ${summary.get('total_msp_spend', 0):,.2f}
MSP Invoices: {summary.get('msp_invoice_count', 0)}
MSP Vendors: {summary.get('msp_vendors_count', 0)}
Services Identified: {summary.get('services_identified', 0)}

MSP Vendor Breakdown:
"""
        
        for msp_vendor, data in msp_services.items():
            total_spend = data.get('total_spend', 0)
            invoice_count = data.get('invoice_count', 0)
            services = data.get('services', {})
            
            section += f"\n{msp_vendor}:\n"
            section += f"• Total Spend: ${total_spend:,.2f}\n"
            section += f"• Invoice Count: {invoice_count}\n"
            
            if services:
                section += "• Services Provided:\n"
                for service, service_data in services.items():
                    service_spend = service_data.get('spend', 0)
                    percentage = (service_spend / total_spend) * 100 if total_spend > 0 else 0
                    section += f"  - {service.replace('_', ' ').title()}: ${service_spend:,.2f} ({percentage:.1f}%)\n"
        
        return section
    
    def create_data_quality_insights(self, data):
        """Create data quality insights section."""
        if 'data_quality' not in data:
            return "Data quality analysis not available."
        
        quality_data = data['data_quality']
        cleaning_analysis = quality_data.get('cleaning_analysis', {})
        quality_metrics = cleaning_analysis.get('quality_metrics', {})
        
        section = f"""
DATA QUALITY INSIGHTS

Original Records: {cleaning_analysis.get('original_count', 0):,}
Final Records: {quality_metrics.get('final_count', 0):,}
Duplicates Removed: {cleaning_analysis.get('duplicates_removed', 0):,}
Data Quality Score: {quality_metrics.get('data_quality_score', 0):.1f}%

Vendor Consolidations: {quality_metrics.get('vendor_consolidation_count', 0)}
Company Consolidations: {quality_metrics.get('company_consolidation_count', 0)}

Consolidation Examples:
"""
        
        vendors_consolidated = cleaning_analysis.get('vendors_consolidated', {})
        for original, consolidated in list(vendors_consolidated.items())[:5]:  # Show first 5
            section += f"• {original} → {consolidated}\n"
        
        return section
    
    def create_cost_optimization_recommendations(self, data):
        """Create cost optimization recommendations section."""
        if 'industry_analysis' not in data:
            return "Industry analysis not available for recommendations."
        
        analysis = data['industry_analysis']
        recommendations = analysis.get('recommendations', [])
        
        section = """
COST OPTIMIZATION RECOMMENDATIONS

Immediate Actions (Next 30 Days):
"""
        
        immediate_actions = [
            "Review contracts with top 3 vendors for negotiation opportunities",
            "Audit license usage to identify unused or underutilized licenses",
            "Implement usage monitoring and alerting systems",
            "Request detailed consumption reports from major vendors"
        ]
        
        for action in immediate_actions:
            section += f"• {action}\n"
        
        section += "\nShort-term Optimizations (Next 90 Days):\n"
        short_term = [
            "Negotiate volume discounts with major vendors",
            "Consider annual payment terms for immediate savings",
            "Implement vendor consolidation strategies",
            "Develop governance policies for license provisioning"
        ]
        
        for optimization in short_term:
            section += f"• {optimization}\n"
        
        section += "\nLong-term Strategies (Next 6-12 Months):\n"
        long_term = [
            "Negotiate Enterprise Agreements for multi-year terms",
            "Implement cloud FinOps practices for ongoing optimization",
            "Develop hybrid cloud strategy to optimize costs",
            "Establish centralized license management system"
        ]
        
        for strategy in long_term:
            section += f"• {strategy}\n"
        
        # Add specific recommendations from analysis
        if recommendations:
            section += "\nSpecific Recommendations from Analysis:\n"
            for i, rec in enumerate(recommendations[:5], 1):  # Top 5
                rec_type = rec.get('type', 'general').replace('_', ' ').title()
                message = rec.get('message', 'No specific message')
                section += f"{i}. {rec_type}: {message}\n"
        
        return section
    
    def create_word_document(self, data, charts):
        """Create the comprehensive Word document."""
        doc = Document()
        
        # Set up document styles
        styles = doc.styles
        title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.size = Pt(18)
        title_style.font.bold = True
        title_style.font.color.rgb = None  # Default color
        
        heading_style = styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
        heading_style.font.size = Pt(14)
        heading_style.font.bold = True
        
        # Title page
        title = doc.add_heading('LICENSING ANALYSIS EXECUTIVE REPORT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph(f'Generated on {datetime.now().strftime("%B %d, %Y")}')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
        
        # Table of Contents placeholder
        toc_heading = doc.add_heading('TABLE OF CONTENTS', level=1)
        toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        toc_items = [
            "Executive Summary",
            "Industry Benchmark Analysis", 
            "MSP Service Analysis",
            "Data Quality Insights",
            "Cost Optimization Recommendations",
            "Appendices"
        ]
        
        for item in toc_items:
            doc.add_paragraph(f"• {item}", style='List Bullet')
        
        doc.add_page_break()
        
        # Executive Summary
        doc.add_heading('EXECUTIVE SUMMARY', level=1)
        summary_text = self.create_executive_summary(data)
        doc.add_paragraph(summary_text)
        
        # Add category spending chart
        if 'category_pie' in charts:
            doc.add_heading('Spending Distribution by Category', level=2)
            doc.add_picture(charts['category_pie'], width=Inches(6))
        
        doc.add_page_break()
        
        # Industry Benchmark Analysis
        doc.add_heading('INDUSTRY BENCHMARK ANALYSIS', level=1)
        benchmark_text = self.create_industry_benchmark_analysis(data)
        doc.add_paragraph(benchmark_text)
        
        # Add vendor chart
        if 'vendors_bar' in charts:
            doc.add_heading('Top Vendors by Spend', level=2)
            doc.add_picture(charts['vendors_bar'], width=Inches(7))
        
        doc.add_page_break()
        
        # MSP Analysis
        doc.add_heading('MSP SERVICE ANALYSIS', level=1)
        msp_text = self.create_msp_analysis(data)
        doc.add_paragraph(msp_text)
        
        # Add company chart
        if 'companies_bar' in charts:
            doc.add_heading('Spending by Company', level=2)
            doc.add_picture(charts['companies_bar'], width=Inches(6))
        
        doc.add_page_break()
        
        # Data Quality Insights
        doc.add_heading('DATA QUALITY INSIGHTS', level=1)
        quality_text = self.create_data_quality_insights(data)
        doc.add_paragraph(quality_text)
        
        doc.add_page_break()
        
        # Cost Optimization Recommendations
        doc.add_heading('COST OPTIMIZATION RECOMMENDATIONS', level=1)
        recommendations_text = self.create_cost_optimization_recommendations(data)
        doc.add_paragraph(recommendations_text)
        
        doc.add_page_break()
        
        # Appendices
        doc.add_heading('APPENDICES', level=1)
        
        # Appendix A: Detailed Vendor Analysis
        doc.add_heading('Appendix A: Detailed Vendor Analysis', level=2)
        if 'cleaned_data' in data:
            vendor_spend = defaultdict(float)
            for item in data['cleaned_data']:
                vendor = item.get('vendor', 'Unknown')
                vendor_spend[vendor] += item.get('total_amount', 0)
            
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Vendor'
            hdr_cells[1].text = 'Total Spend'
            hdr_cells[2].text = 'Percentage'
            
            total_spend = sum(vendor_spend.values())
            for vendor, spend in sorted(vendor_spend.items(), key=lambda x: x[1], reverse=True):
                row_cells = table.add_row().cells
                row_cells[0].text = vendor
                row_cells[1].text = f"${spend:,.2f}"
                row_cells[2].text = f"{(spend/total_spend)*100:.1f}%"
        
        doc.add_page_break()
        
        # Appendix B: Methodology
        doc.add_heading('Appendix B: Methodology', level=2)
        methodology_text = """
This analysis was conducted using the following methodology:

1. Data Collection: Invoice data was collected from multiple sources and consolidated
2. Data Cleaning: Duplicates were removed and vendor names were standardized
3. Vendor Consolidation: Intelligent algorithms consolidated vendor name variations
4. Company Extraction: Company names were extracted from billing information
5. Industry Benchmarking: Spending was compared against industry standards
6. MSP Analysis: Managed Service Provider services were broken down by underlying services
7. Cost Optimization: Recommendations were generated based on analysis findings

The analysis covers the period from 2024-2025 and includes all licensing-related expenditures.
"""
        doc.add_paragraph(methodology_text)
        
        return doc
    
    def generate_report(self):
        """Generate the comprehensive executive report."""
        print("Creating Comprehensive Executive Report...")
        print("Loading analysis data...")
        
        # Load data
        data = self.load_data()
        
        if not data:
            print("No data available for report generation!")
            return
        
        print("Creating professional charts...")
        # Create charts
        charts = self.create_charts(data)
        
        print("Generating Word document...")
        # Create Word document
        doc = self.create_word_document(data, charts)
        
        # Ensure output directory exists
        os.makedirs(os.path.dirname(self.output_file), exist_ok=True)
        
        # Save document
        doc.save(self.output_file)
        
        print(f"Comprehensive executive report completed!")
        print(f"Report saved to: {self.output_file}")
        print(f"Charts created: {len(charts)}")
        
        # Print summary
        if 'cleaned_data' in data:
            total_spend = sum(item.get('total_amount', 0) for item in data['cleaned_data'])
            print(f"Total spend analyzed: ${total_spend:,.2f}")
            print(f"Records analyzed: {len(data['cleaned_data']):,}")

def main():
    """Main execution function."""
    generator = ComprehensiveExecutiveReport()
    generator.generate_report()

if __name__ == "__main__":
    main() 