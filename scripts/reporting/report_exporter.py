#!/usr/bin/env python3
"""
Report Exporter
Converts markdown reports to PDF and Word documents with charts
"""

import json
import os
from datetime import datetime
from collections import defaultdict
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import pandas as pd

class ReportExporter:
    def __init__(self):
        self.report_date = datetime.now().strftime("%B %d, %Y")
        
    def create_charts(self, analysis_results, categorized_spend):
        """Create charts for the report."""
        charts_dir = "reports/charts"
        os.makedirs(charts_dir, exist_ok=True)
        
        # Set style
        plt.style.use('default')
        plt.rcParams['figure.figsize'] = (12, 8)
        plt.rcParams['font.size'] = 10
        
        # 1. Spending by Category (Pie Chart)
        self.create_spending_pie_chart(analysis_results, charts_dir)
        
        # 2. Industry Benchmark Comparison (Bar Chart)
        self.create_benchmark_comparison_chart(analysis_results, charts_dir)
        
        # 3. Cost Variance Analysis (Horizontal Bar Chart)
        self.create_variance_chart(analysis_results, charts_dir)
        
        # 4. Vendor Distribution (Bar Chart)
        self.create_vendor_distribution_chart(categorized_spend, charts_dir)
        
        # 5. Monthly Spending Trend (Line Chart)
        self.create_monthly_trend_chart(categorized_spend, charts_dir)
        
        return charts_dir
    
    def create_spending_pie_chart(self, analysis_results, charts_dir):
        """Create pie chart showing spending by category."""
        categories = []
        amounts = []
        colors = ['#FF6B6B', '#4ECDC4', '#45B7D1', '#96CEB4', '#FFEAA7']
        
        for category, data in analysis_results.items():
            categories.append(category.replace('_', ' ').title())
            amounts.append(data['total_spend'])
        
        fig, ax = plt.subplots(figsize=(10, 8))
        wedges, texts, autotexts = ax.pie(amounts, labels=categories, autopct='%1.1f%%', 
                                         colors=colors[:len(categories)], startangle=90)
        
        ax.set_title('Spending Distribution by Category', fontsize=16, fontweight='bold', pad=20)
        
        # Add total amount in center
        total = sum(amounts)
        ax.text(0, 0, f'Total: ${total:,.0f}', ha='center', va='center', 
                fontsize=14, fontweight='bold')
        
        plt.tight_layout()
        plt.savefig(f'{charts_dir}/spending_by_category.png', dpi=300, bbox_inches='tight')
        plt.close()
    
    def create_benchmark_comparison_chart(self, analysis_results, charts_dir):
        """Create bar chart comparing actual vs industry benchmark."""
        categories = []
        actual_spend = []
        benchmark_low = []
        benchmark_high = []
        
        for category, data in analysis_results.items():
            categories.append(category.replace('_', ' ').title())
            actual_spend.append(data['total_spend'])
            benchmark_low.append(data['benchmark']['low'])
            benchmark_high.append(data['benchmark']['high'])
        
        x = np.arange(len(categories))
        width = 0.35
        
        fig, ax = plt.subplots(figsize=(12, 8))
        
        # Create bars
        bars1 = ax.bar(x - width/2, actual_spend, width, label='Actual Spend', color='#FF6B6B')
        bars2 = ax.bar(x + width/2, benchmark_high, width, label='Industry Benchmark (High)', color='#4ECDC4')
        
        # Add benchmark low as error bars
        ax.errorbar(x + width/2, benchmark_high, yerr=[np.array(benchmark_high) - np.array(benchmark_low)], 
                   fmt='none', color='black', capsize=5)
        
        ax.set_xlabel('Categories', fontsize=12)
        ax.set_ylabel('Spend Amount ($)', fontsize=12)
        ax.set_title('Actual Spend vs Industry Benchmark', fontsize=16, fontweight='bold')
        ax.set_xticks(x)
        ax.set_xticklabels(categories, rotation=45, ha='right')
        ax.legend()
        
        # Add value labels on bars
        for bar in bars1:
            height = bar.get_height()
            ax.text(bar.get_x() + bar.get_width()/2., height,
                   f'${height:,.0f}', ha='center', va='bottom', fontsize=8)
        
        plt.tight_layout()
        plt.savefig(f'{charts_dir}/benchmark_comparison.png', dpi=300, bbox_inches='tight')
        plt.close()
    
    def create_variance_chart(self, analysis_results, charts_dir):
        """Create horizontal bar chart showing cost variances."""
        categories = []
        variances = []
        colors = []
        
        for category, data in analysis_results.items():
            categories.append(category.replace('_', ' ').title())
            variance = data['variance']['percentage']
            variances.append(variance)
            
            # Color coding based on assessment
            if data['variance']['assessment'] == 'above_industry_standard':
                colors.append('#FF6B6B')  # Red for above
            elif data['variance']['assessment'] == 'below_industry_standard':
                colors.append('#4ECDC4')  # Green for below
            else:
                colors.append('#FFEAA7')  # Yellow for within
        
        fig, ax = plt.subplots(figsize=(10, 6))
        
        bars = ax.barh(categories, variances, color=colors)
        ax.set_xlabel('Variance from Industry Standard (%)', fontsize=12)
        ax.set_title('Cost Variance Analysis', fontsize=16, fontweight='bold')
        
        # Add vertical line at 0
        ax.axvline(x=0, color='black', linestyle='-', alpha=0.3)
        
        # Add value labels
        for i, bar in enumerate(bars):
            width = bar.get_width()
            ax.text(width, bar.get_y() + bar.get_height()/2, 
                   f'{width:+.1f}%', ha='left' if width > 0 else 'right', 
                   va='center', fontsize=10, fontweight='bold')
        
        plt.tight_layout()
        plt.savefig(f'{charts_dir}/cost_variance.png', dpi=300, bbox_inches='tight')
        plt.close()
    
    def create_vendor_distribution_chart(self, categorized_spend, charts_dir):
        """Create bar chart showing vendor distribution."""
        vendor_counts = defaultdict(int)
        
        for category, items in categorized_spend.items():
            for item in items:
                vendor = item.get('vendor', 'Unknown')
                vendor_counts[vendor] += 1
        
        # Get top 10 vendors
        top_vendors = sorted(vendor_counts.items(), key=lambda x: x[1], reverse=True)[:10]
        vendors, counts = zip(*top_vendors)
        
        fig, ax = plt.subplots(figsize=(12, 8))
        
        bars = ax.bar(range(len(vendors)), counts, color='#45B7D1')
        ax.set_xlabel('Vendors', fontsize=12)
        ax.set_ylabel('Number of Invoices', fontsize=12)
        ax.set_title('Vendor Distribution (Top 10)', fontsize=16, fontweight='bold')
        ax.set_xticks(range(len(vendors)))
        ax.set_xticklabels(vendors, rotation=45, ha='right')
        
        # Add value labels
        for bar in bars:
            height = bar.get_height()
            ax.text(bar.get_x() + bar.get_width()/2., height,
                   f'{int(height)}', ha='center', va='bottom', fontsize=10)
        
        plt.tight_layout()
        plt.savefig(f'{charts_dir}/vendor_distribution.png', dpi=300, bbox_inches='tight')
        plt.close()
    
    def create_monthly_trend_chart(self, categorized_spend, charts_dir):
        """Create line chart showing monthly spending trends."""
        monthly_data = defaultdict(float)
        
        for category, items in categorized_spend.items():
            for item in items:
                date_str = item.get('invoice_date', '')
                if 'July' in date_str:
                    monthly_data['July 2025'] += item.get('total_amount', 0)
                elif 'June' in date_str:
                    monthly_data['June 2025'] += item.get('total_amount', 0)
                elif 'May' in date_str:
                    monthly_data['May 2025'] += item.get('total_amount', 0)
        
        if monthly_data:
            months = list(monthly_data.keys())
            amounts = list(monthly_data.values())
            
            fig, ax = plt.subplots(figsize=(10, 6))
            
            ax.plot(months, amounts, marker='o', linewidth=3, markersize=8, color='#4ECDC4')
            ax.set_xlabel('Month', fontsize=12)
            ax.set_ylabel('Total Spend ($)', fontsize=12)
            ax.set_title('Monthly Spending Trend', fontsize=16, fontweight='bold')
            
            # Add value labels
            for i, amount in enumerate(amounts):
                ax.text(i, amount, f'${amount:,.0f}', ha='center', va='bottom', fontsize=10)
            
            plt.tight_layout()
            plt.savefig(f'{charts_dir}/monthly_trend.png', dpi=300, bbox_inches='tight')
            plt.close()
    
    def export_to_word(self, analysis_results, categorized_spend, charts_dir):
        """Export analysis to Word document."""
        doc = Document()
        
        # Title
        title = doc.add_heading('Licensing Cost Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle = doc.add_paragraph(f'Industry Benchmark Analysis - {self.report_date}')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=1)
        total_spend = sum(result['total_spend'] for result in analysis_results.values())
        above_standard = sum(1 for result in analysis_results.values() 
                           if result['variance']['assessment'] == 'above_industry_standard')
        
        summary_text = f"""
This report analyzes {len(analysis_results)} spending categories with a total spend of ${total_spend:,.2f}. 
{above_standard} out of {len(analysis_results)} categories are above industry standards, 
indicating significant optimization opportunities.
        """
        doc.add_paragraph(summary_text.strip())
        
        # Add spending pie chart
        doc.add_heading('Spending Distribution by Category', level=2)
        doc.add_picture(f'{charts_dir}/spending_by_category.png', width=Inches(6))
        
        # Category Analysis
        doc.add_heading('Detailed Category Analysis', level=1)
        
        for category, data in analysis_results.items():
            doc.add_heading(f'{category.replace("_", " ").title()}', level=2)
            
            # Category details
            details = f"""
• Total Spend: ${data['total_spend']:,.2f}
• Percentage of Total: {data['percentage_of_total']:.1f}%
• Invoice Count: {data['invoice_count']}
• Industry Benchmark: ${data['benchmark']['low']:,.2f} - ${data['benchmark']['high']:,.2f}
• Assessment: {self.get_assessment_description(data['variance']['assessment'])}
• Variance: {data['variance']['percentage']:+.1f}% from industry standard
            """
            doc.add_paragraph(details.strip())
        
        # Add benchmark comparison chart
        doc.add_heading('Industry Benchmark Comparison', level=2)
        doc.add_picture(f'{charts_dir}/benchmark_comparison.png', width=Inches(6))
        
        # Cost variance chart
        doc.add_heading('Cost Variance Analysis', level=2)
        doc.add_picture(f'{charts_dir}/cost_variance.png', width=Inches(6))
        
        # Vendor distribution
        doc.add_heading('Vendor Distribution', level=2)
        doc.add_picture(f'{charts_dir}/vendor_distribution.png', width=Inches(6))
        
        # Recommendations
        doc.add_heading('Recommendations', level=1)
        
        optimization_opportunities = []
        for category, data in analysis_results.items():
            if data['variance']['assessment'] == 'above_industry_standard':
                potential_savings = data['total_spend'] * (data['variance']['percentage'] / 100)
                optimization_opportunities.append({
                    'category': category,
                    'potential_savings': potential_savings
                })
        
        total_potential_savings = sum(opp['potential_savings'] for opp in optimization_opportunities)
        
        recommendations = f"""
Total Potential Savings: ${total_potential_savings:,.2f}

Immediate Actions (Next 30 Days):
1. Focus on categories above industry standards
2. Negotiate with high-variance vendors
3. Review vendor contracts for better pricing

Strategic Actions (Next 90 Days):
1. Implement vendor diversification
2. Establish category-specific benchmarks
3. Optimize spending allocation

Long-term Strategy (Next 12 Months):
1. Develop category management
2. Implement automated benchmarking
3. Strategic vendor partnerships
        """
        doc.add_paragraph(recommendations.strip())
        
        # Save document
        reports_dir = "reports/executive"
        os.makedirs(reports_dir, exist_ok=True)
        output_file = f"{reports_dir}/licensing_analysis_report_{datetime.now().strftime('%Y%m%d')}.docx"
        doc.save(output_file)
        
        return output_file
    
    def get_assessment_description(self, assessment):
        """Get human-readable assessment description."""
        descriptions = {
            'below_industry_standard': 'Below Industry Standard (Good)',
            'within_industry_standard': 'Within Industry Standard (Good)',
            'above_industry_standard': 'Above Industry Standard (Needs Attention)',
            'unknown': 'Unknown (Needs Investigation)'
        }
        return descriptions.get(assessment, assessment)

def main():
    """Generate charts and export reports."""
    
    # Load the data
    data_file = "reports/processed_licensing_data.json"
    if not os.path.exists(data_file):
        print(f"❌ Data file not found: {data_file}")
        return
    
    with open(data_file, 'r') as f:
        data = json.load(f)
    
    print(f"Loaded {len(data)} records for report generation")
    
    # Import the analyzer
    from industry_benchmark_analyzer import IndustryBenchmarkAnalyzer
    
    # Initialize analyzer and exporter
    analyzer = IndustryBenchmarkAnalyzer()
    exporter = ReportExporter()
    
    # Perform analysis
    print("Analyzing industry benchmarks...")
    analysis_results, categorized_spend = analyzer.analyze_industry_comparison(data)
    
    # Create charts
    print("Creating charts...")
    charts_dir = exporter.create_charts(analysis_results, categorized_spend)
    
    # Export to Word
    print("Exporting to Word document...")
    word_file = exporter.export_to_word(analysis_results, categorized_spend, charts_dir)
    
    print(f"Word document saved to: {word_file}")
    print(f"Charts saved to: {charts_dir}/")
    
    print(f"\nGenerated Files:")
    print(f"   - Word Document: {word_file}")
    print(f"   - Charts Directory: {charts_dir}/")
    print(f"   - 5 Professional Charts Created")

if __name__ == "__main__":
    main() 